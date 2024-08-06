import os
import requests
import streamlit as st
from io import BytesIO
from docx import Document
from langchain_groq import ChatGroq
from langchain_openai import ChatOpenAI
from langchain_google_genai import ChatGoogleGenerativeAI
from crewai import Agent, Task, Crew, Process
from crewai_tools import ScrapeWebsiteTool, SerperDevTool
import asyncio
from datetime import datetime

def verify_api_key(api_url, headers):
    try:
        response = requests.get(api_url, headers=headers)
        response.raise_for_status()  # Raises an HTTPError for bad responses
        return True
    except requests.exceptions.HTTPError:
        return False
    except requests.exceptions.RequestException as e:
        raise ValueError(f"An error occurred: {str(e)}")

def verify_gemini_api_key(api_key):
    api_url = f"https://generativelanguage.googleapis.com/v1/models?key={api_key}"
    headers = {'Content-Type': 'application/json'}
    return verify_api_key(api_url, headers)

def verify_gpt_api_key(api_key):
    api_url = "https://api.openai.com/v1/models"
    headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
    return verify_api_key(api_url, headers)

def verify_groq_api_key(api_key):
    api_url = "https://api.groq.com/openai/v1/models"
    headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
    return verify_api_key(api_url, headers)

def verify_serperdev_api_key(api_key):
    api_url = "https://google.serper.dev/search"
    headers = {"X-API-KEY": api_key, "Content-Type": "application/json"}
    payload = {"q": "test"}
    try:
        response = requests.post(api_url, json=payload, headers=headers)
        response.raise_for_status()
        return True
    except requests.exceptions.HTTPError:
        return False
    except requests.exceptions.RequestException as e:
        raise ValueError(f"An error occurred: {str(e)}")


def generate_event_plan(llm, serper_dev_key,event_type, budget, num_people, cuisine, date, location):
    os.environ['SERPER_API_KEY'] = serper_dev_key
    
    search_tool = SerperDevTool(location=location, n_results=5)
    scrape_tool = ScrapeWebsiteTool()

    def create_agent(role, goal, backstory, llm):
        return Agent(
            role=role,
            goal=goal,
            backstory=backstory,
            verbose=True,
            allow_delegation=False,
            llm=llm
        )

    venue_coordinator = create_agent(
        role='Venue Coordinator',
        goal='Select and secure a specific, suitable venue for the event at the specified location',
        backstory="An expert in venue selection with extensive knowledge of event spaces in the given location. You have connections with various venues and can provide detailed information about their offerings.",
        llm=llm
    )

    catering_manager = create_agent(
        role='Catering Manager',
        goal='Create a detailed menu plan with specific dishes for the event',
        backstory="A culinary expert with experience in menu planning and dietary accommodations. You have in-depth knowledge of the specified cuisine and can suggest popular local catering services.",
        llm=llm
    )

    entertainment_coordinator = create_agent(
        role='Entertainment Coordinator',
        goal='Plan specific entertainment and activities for the event',
        backstory="A creative professional with connections in the local entertainment industry. You're familiar with popular and unique entertainment options suitable for various event types.",
        llm=llm
    )

    logistics_coordinator = create_agent(
        role='Logistics Coordinator',
        goal='Plan all logistical aspects of the event with specific details',
        backstory="An organized professional with experience in coordinating event logistics. You have knowledge of local transportation options, equipment rental services, and can create detailed schedules.",
        llm=llm
    )

    event_presenter = create_agent(
        role='Presentation Coordinator',
        goal='Prepare a comprehensive and specific event plan document',
        backstory="An expert in creating detailed event plans and presentations. You excel at organizing information and presenting it in a clear, structured format.",
        llm=llm
    )

    task_venue = Task(
    description=f"""Select a specific venue for the {event_type} event on {date} in {location} accommodating {num_people} people within the budget of {budget}. Provide detailed information including:
    1. Venue name and full address
    2. Capacity (ensure it comfortably fits {num_people})
    3. Exact rental cost and any additional fees
    4. Complete list of amenities which will be useful for the event {event_type} (e.g., A/V equipment, Wi-Fi, parking, accessibility features)
    5. Confirm availability for {date} and any specific time restrictions
    6. Vendor policies (e.g., preferred caterers, outside vendor restrictions)
    7. Insurance requirements
    8. Setup and breakdown times allowed
    9. Nearby accommodations for out-of-town guests
    
    The venue should be appropriate for the event type and location. Use the web search tool and web scraper to get real, up-to-date information about venues in {location}. Contact the venue directly if possible to confirm details.""",
    agent=venue_coordinator,
    expected_output="A comprehensive description of the selected venue with all requested details, ensuring it meets the event requirements and budget constraints.",
    tools=[search_tool, scrape_tool],
    async_execution = True
)

    task_entertainment = Task(
    description=f"""Plan 3-4 specific entertainment options suitable for a {event_type} event with {num_people} attendees on {date} in {location}. For each option, provide:
    1. Type of entertainment (e.g., live band, DJ, interactive performance, team-building activity)
    2. Specific performer/company name and their background
    3. Duration of performance/activity
    4. Detailed cost breakdown (performance fee, equipment rental, travel expenses)
    5. How it fits the event theme and engages the audience
    6. Technical requirements (stage size, sound system, lighting)
    7. Any special accommodations needed
    8. Past client testimonials or reviews
    9. Backup plans in case of cancellation
    10. Suggestions for event flow (e.g., when to schedule each entertainment option)
    Use the web search tool and web scraper to find popular and appropriate entertainment options in {location}. Contact entertainment companies or performers directly if possible to confirm availability and get accurate quotes.""",
    agent=entertainment_coordinator,
    expected_output="A detailed list of 3-4 entertainment options including all requested information, ensuring they are well-suited to the event type, budget, and location.",
    tools=[search_tool, scrape_tool],
    async_execution = True
)

    task_logistics = Task(
    description=f"""Create a detailed logistics plan for the {event_type} event on {date} in {location} for {num_people} people. Include:
    1. Transportation:
       - Shuttle service options from major hotels or nearby airports
       - Parking arrangements and valet services if needed
       - Public transportation routes and schedules with links
    2. Equipment rentals:
       - Detailed list of required A/V equipment (microphones, speakers, projectors, screens)
       - Furniture needs (tables, chairs, linens, decorative items)
       - Any specialized equipment for activities or presentations
    3. Staffing:
       - Number and roles of staff needed (greeters, servers, technicians, security)
       - Recommendations for reputable staffing agencies in {location}
    4. Detailed event schedule:
       - Setup time and process
       - Registration/check-in procedure
       - Minute-by-minute timeline of activities
       - Meal and break times
       - Entertainment slots
       - Networking opportunities
       - Cleanup and breakdown process
    5. Safety and security measures:
       - First aid stations
       - Emergency protocols
       - Security personnel requirements
    6. Signage and wayfinding plan
    7. Contingency plans for weather (if applicable) or other potential issues
    8. Sustainability considerations (waste management, eco-friendly options)
    9. Any necessary permits or licenses required for the event
    10. Estimated costs for each aspect of the logistics plan
    Use the web search tool and web scraper to find reputable local transportation, equipment rental, and staffing services in {location}. Contact service providers directly if possible to get accurate quotes and confirm availability.""",
    agent=logistics_coordinator,
    expected_output="A comprehensive logistics plan covering all requested aspects, with specific details, service provider recommendations, and accurate cost estimates.",
    tools=[search_tool, scrape_tool],
    async_execution = True
)
    task_catering = Task(
    description=f"""Create a detailed menu plan for {num_people} people featuring {cuisine} cuisine for the event on {date} in {location}. Include:
    1. 4-5 options each for appetizers, main courses, desserts, and beverages
    2. Specific dish names and brief descriptions
    3. Accommodations for common dietary restrictions (vegan, vegetarian, gluten-free, nut-free, kosher, halal)
    4. Estimated cost per person and total catering cost
    5. Suggest 2-3 highly-rated local catering services that specialize in {cuisine}
    6. For each caterer, provide:
       - Company name and contact information
       - Years of experience and notable past events
       - Specific services offered (e.g., staff, equipment rental, setup/cleanup)
       - Customer reviews and ratings
    7. Bar service options (if applicable) and estimated costs
    8. Any special considerations for {event_type} events
    Use the web search tool and web scraper to find popular {cuisine} dishes, local catering services, and customer reviews in {location}. If possible, contact the caterers for quotes and menu customization options.""",
    agent=catering_manager,
    expected_output="A comprehensive menu plan with all requested details, including specific dishes, costs, and thoroughly researched local catering options.",
    tools=[search_tool, scrape_tool],
    async_execution = True
)

    event_presentation = Task(
        description="Compile all the information from the venue, catering, entertainment, and logistics tasks into a comprehensive and coherent event plan document. Organize the information in a clear, logical structure. Include an event overview, venue details, menu plan, entertainment options, logistics plan, and a budget breakdown. Ensure all specific details are included and any placeholders are removed.",
        agent=event_presenter,
        expected_output="A detailed, well-structured event plan document covering  all aspects of the event with specific information and no placeholders.",
        context=[task_venue, task_catering, task_entertainment, task_logistics],
    )

    crew = Crew(
        agents=[venue_coordinator, catering_manager, entertainment_coordinator, logistics_coordinator, event_presenter],
        tasks=[task_venue, task_catering, task_entertainment, task_logistics, event_presentation],
        verbose=2,
        manager_agent=Agent(
            role='Event Planning Manager',
            goal=f'Oversee and coordinate the entire event planning process for the {event_type} in {location} on {date} for {num_people} people, ensuring all aspects are within the total budget of {budget} in the local currency of {location}. Please utilize the entire budget and not be cheap',
            backstory="An experienced event planner with a keen eye for detail and excellent organizational skills. You have extensive knowledge of event planning in various locations and can make informed decisions to create a cohesive and successful event.",
            verbose=True,
            allow_delegation=True,
            llm=llm
        ),
        process=Process.hierarchical,

    )

    result = crew.kickoff()
    
    return result

def main():
    st.header('AI Event Planner')
    validity_model = False
    validity_serperdev = False

    # Initialize session state
    if 'generated_content' not in st.session_state:
        st.session_state.generated_content = None

    with st.sidebar:
        with st.form('Gemini/OpenAI/Groq'):
            model = st.radio('Choose Your LLM', ('Gemini', 'OpenAI', 'Groq'))
            api_key = st.text_input(f'Enter your API key', type="password")
            serper_dev_key = st.text_input(f'Enter your SerperDev API key', type="password")
            submitted = st.form_submit_button("Submit")
            
        if api_key and serper_dev_key:
            if model == "Gemini":
                validity_model = verify_gemini_api_key(api_key)
            elif model == "OpenAI":
                validity_model = verify_gpt_api_key(api_key)
            elif model == "Groq":
                validity_model = verify_groq_api_key(api_key)
            
            validity_serperdev = verify_serperdev_api_key(serper_dev_key)
            
            if validity_model:
                st.write(f"Valid {model} API key")
            else:
                st.write(f"Invalid {model} API key")
            
            if validity_serperdev:
                st.write("Valid SerperDev API key")
            else:
                st.write("Invalid SerperDev API key")

    if validity_model and validity_serperdev:
        if model == 'OpenAI':
            async def setup_OpenAI():
                loop = asyncio.get_event_loop()
                if loop is None:
                    loop = asyncio.new_event_loop()
                    asyncio.set_event_loop(loop)

                llm = ChatOpenAI(model='gpt-4-turbo', temperature=0.6, max_tokens=10000, api_key=api_key)
                return llm

            llm = asyncio.run(setup_OpenAI())

        elif model == 'Gemini':
            async def setup_gemini():
                loop = asyncio.get_event_loop()
                if loop is None:
                    loop = asyncio.new_event_loop()
                    asyncio.set_event_loop(loop)

                llm = ChatGoogleGenerativeAI(
                    model="gemini-1.5-flash",
                    verbose=True,
                    temperature=0.6,
                    google_api_key=api_key
                )
                return llm

            llm = asyncio.run(setup_gemini())

        elif model == 'Groq':
            async def setup_groq():
                loop = asyncio.get_event_loop()
                if loop is None:
                    loop = asyncio.new_event_loop()
                    asyncio.set_event_loop(loop)

                llm = ChatGroq(
                    api_key=api_key,
                    model='llama-3.1-70b-versatile',
                    max_tokens=5000,
                    temperature=0.6
                )
                return llm

            llm = asyncio.run(setup_groq())

        event = st.text_input("Enter your event type (Example - Conference for AI)")
        location = st.text_input("Enter location where event needs to be planned. Eg- Pune,India")
        budget = st.number_input("Enter your budget", min_value=0, step=2000)
        num_people = st.number_input("Enter total number of people", min_value=1, step=1)
        cuisine = st.text_input("Enter food cuisine needed")
        date = st.date_input("Select event date", datetime.now().date())

        if st.button("Generate Event Plan"):
            with st.spinner("Generating event plan..."):
                st.session_state.generated_content = generate_event_plan(llm, serper_dev_key,event, budget, num_people, cuisine, date, location)

        if st.session_state.generated_content:
            st.markdown(st.session_state.generated_content)

            doc = Document()
            doc.add_heading(f"{event} - {date}", 0)
            doc.add_paragraph(st.session_state.generated_content)

            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            st.download_button(
                label="Download as Word Document",
                data=buffer,
                file_name=f"{event}_{date}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

if __name__ == "__main__":
    main()
